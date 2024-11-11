import os
import pandas as pd
import random
from docx import Document

class_id_to_time_map = {
    "A1501": "1(56)", "A1520": "1(78)", "A1521": "2(34)", "A1507": "2(56)", "A1509": "2(78)",
    "A1536": "3(34)", "A1525": "3(56)", "A1510": "3(78)", "A1512": "4(34)", "A1513": "4(56)",
    "A1535": "4(78)", "A1515": "5(34)", "A1517": "5(56)"
}

# 假設每個班級的座位表格式不同
# 我們會用一個字典來指定每個班級的座位行數和列數
class_seating_layout = {
    'A1501': {'rows': 7, 'cols': 7},
    'A1520': {'rows': 7, 'cols': 7},
    'A1521': {'rows': 6, 'cols': 9},
    'A1507': {'rows': 6, 'cols': 9},
    'A1509': {'rows': 6, 'cols': 9},
    'A1536': {'rows': 5, 'cols': 10},
    'A1525': {'rows': 7, 'cols': 9},
    'A1510': {'rows': 7, 'cols': 7},
    'A1512': {'rows': 8, 'cols': 7},
    'A1513': {'rows': 6, 'cols': 8},
    'A1535': {'rows': 7, 'cols': 8},
    'A1515': {'rows': 5, 'cols': 10},
    'A1517': {'rows': 5, 'cols': 10}
    # 可以繼續增加其他班級
}

import pandas as pd

def generate_seating_chart(students_df, class_name):
    """
    根據學生名單生成座位表，並隨機安排座位，返回 DataFrame 格式
    """
    if class_name not in class_seating_layout:
        return None  # 如果班級名稱不在座位配置中，返回 None

    seating_layout = class_seating_layout[class_name]
    rows = seating_layout['rows']
    cols = seating_layout['cols']

    # 隨機打亂學生順序並去掉第一列
    students_df = students_df.drop(0).sample(frac=1).reset_index(drop=True)
    # print(students_df.head())
    # 如果學生數量超過座位數，則只保留需要的學生數量
    max_seats = rows * cols
    if len(students_df) > max_seats:
        students_df = students_df.head(max_seats)
    
    # 添加空座位以匹配座位數量（若學生數量少於座位數，補上空白）
    # empty_seats_needed = max_seats - len(students_df)
    # if empty_seats_needed > 0:
    #     empty_seats_df = pd.DataFrame([{'座號': '', '系 年 班': '', '姓名(Name)': ''}] * empty_seats_needed)
    #     students_df = pd.concat([students_df, empty_seats_df], ignore_index=True)

    # 重整座位表為指定的 rows 和 cols 結構
    # seating_chart = students_df.values.reshape(rows, cols, -1)  # 轉換成三維陣列

    # 將座位表轉回 DataFrame 格式
    # seating_chart_df = pd.DataFrame(seating_chart.reshape(rows * cols, -1), columns=['座號', '系 年 班', '姓名(Name)'])
    seating_chart_df = pd.DataFrame(students_df, columns=['座號', '系 年 班', '姓名(Name)'])
    # print(seating_chart_df.head())
    return seating_chart_df


from docx import Document

def save_seating_chart_to_word(seating_chart, class_name, file_path):
    """
    將座位表儲存為 Word 檔案，座號、系級和姓名填寫在單一格內。
    """
    doc = Document()
    doc.add_heading(f"{class_id_to_time_map[class_name]} 座位表", 0)

    # 取得行列數
    rows = class_seating_layout[class_name]['rows']
    cols = class_seating_layout[class_name]['cols']
    
    # 插入表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for i in range(rows):
        for j in range(cols):
            try:
                if i*cols+j >= len(seating_chart):
                    break
                student = seating_chart.iloc[i*cols+j]
                table.cell(i, j).text = f"{student['座號']}\n{student['系 年 班']}\n{student['姓名(Name)']}"
                # print(table.cell(i, j).text)
            except StopIteration:
                # 若學生不足，剩餘格子留空白
                table.cell(i, j).text = ""

    # 儲存 Word 檔案至完整路徑
    doc.save(file_path)
    print(f"已儲存座位表至 {file_path}")

def generate_seating_chart_for_all_csvs(root_folder,word_folder,result_callback):
    """
    遍歷資料夾中的 CSV 檔案，並為每個檔案生成座位表
    """
    results = []
    try:
        for dirpath, dirnames, filenames in os.walk(root_folder):
            for filename in filenames:
                if filename.endswith('.csv'):
                    csv_path = os.path.join(dirpath, filename)
                    # 讀取 CSV 檔案
                    df = pd.read_csv(csv_path)

                    # 篩選出成績欄位為空白的學生
                    students_to_seat = df[df['成績(Score)'].isnull()]
                    if not students_to_seat.empty:
                        # 根據班級名稱生成座位表
                        folder_name = os.path.basename(dirpath)
                        class_name = f"{folder_name}_{os.path.splitext(filename)[0]}"[:5]  # 使用檔案名作為班級名稱  
                        result_callback("開始處理班級: " + class_name)

                        seating_chart = generate_seating_chart(students_to_seat, class_name)
                        result_callback("座位表生成完成")
                        try:
                            doc_name = f'{class_id_to_time_map[class_name]}_考試座位表.docx'
                            word_folder_path = os.path.join(word_folder, doc_name)
                            save_seating_chart_to_word(seating_chart, class_name, word_folder_path)
                            results.append(f"成功產生 {class_id_to_time_map[class_name]} 座位表並儲存為 {word_folder_path}")
                        except Exception as e:
                            results.append(f"班級 {class_name} 的座位表生成失敗: {e}")
                    else:
                        results.append(f"班級 {filename} 沒有需要填入座位表的學生")
    except Exception as e:
        results.append(f"處理資料夾時發生錯誤: {e}")
    return results
